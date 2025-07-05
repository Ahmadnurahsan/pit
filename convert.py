import requests
from docx import Document
from docx.shared import Inches

def get_data_from_api(api_url):
    """Fungsi untuk mengambil data dari URL API."""
    try:
        response = requests.get(api_url)
        response.raise_for_status()  # Cek jika ada error HTTP
        return response.json()
    except requests.exceptions.RequestException as e:
        print(f"Gagal mengambil data dari API: {e}")
        return None

def create_docx_from_json(data, filename="dokumen_output.docx"):
    """
    Membuat file .docx dari data JSON.
    Data diharapkan berupa list of dictionaries dengan key 'title' dan 'content'.
    """
    # 1. Buat dokumen baru di memori
    document = Document()

    # 2. Tambahkan judul utama dokumen (opsional)
    # level=0 adalah untuk gaya Judul Utama (Title) di Word
    document.add_heading(filename.replace('.docx', '').replace('_', ' ').title(), level=0)

    # 3. Looping melalui data dan tambahkan ke dokumen
    for item in data:
        title = item.get('title', 'Tanpa Judul')
        content = item.get('content', 'Tidak ada konten.')

        # Tambahkan judul item sebagai Heading 1
        document.add_heading(title, level=1)

        # Tambahkan konten sebagai paragraf biasa
        document.add_paragraph(content)

        # Tambahkan spasi antar item (opsional)
        document.add_paragraph()

    # 4. Simpan dokumen ke file lokal
    try:
        document.save(filename)
        print(f"Sukses! File '{filename}' telah berhasil dibuat. 💾")
    except Exception as e:
        print(f"Gagal menyimpan file: {e}")


# --- CONTOH PENGGUNAAN ---

if __name__ == '__main__':
    # Opsi 1: Menggunakan data JSON statis yang ada di dalam kode
    sample_json_data = [
        {
            "title": "Laporan Penjualan Kuartal Pertama",
            "content": "Penjualan pada Q1 mengalami peningkatan signifikan sebesar 20% year-over-year."
        },
        {
            "title": "Analisis Produk",
            "content": "Produk seri X menjadi kontributor utama dengan margin keuntungan tertinggi."
        }
    ]

    print("Membuat dokumen dari data statis...")
    create_docx_from_json(sample_json_data, "Laporan_Penjualan.docx")

    # # Opsi 2: Mengambil data dari API (contoh: JSONPlaceholder)
    # print("\nMembuat dokumen dari data API...")
    # posts_api_url = "https://jsonplaceholder.typicode.com/posts?_limit=3"
    # api_data = get_data_from_api(posts_api_url)
    
    # if api_data:
    #     # Sesuaikan format data dari API jika perlu
    #     formatted_data = [{"title": post['title'], "content": post['body']} for post in api_data]
    #     create_docx_from_json(formatted_data, "Laporan_Dari_API.docx")
